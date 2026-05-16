"""
File utility functions for text extraction from different file formats
"""
import os
from typing import Optional
from docx import Document as DocxDocument
from PyPDF2 import PdfReader


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from .txt file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from .docx file"""
    doc = DocxDocument(file_path)
    text_parts = []
    
    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    
    return '\n'.join(text_parts)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from .pdf file"""
    text_parts = []
    
    with open(file_path, 'rb') as f:
        pdf_reader = PdfReader(f)
        
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
    
    return '\n'.join(text_parts)


def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """
    Extract text from file based on extension.
    
    Args:
        file_path: Path to the file
        file_extension: File extension (.txt, .docx, .pdf)
    
    Returns:
        Extracted text content
    
    Raises:
        ValueError: If file type is not supported
    """
    file_extension = file_extension.lower()
    
    if file_extension == '.txt':
        return extract_text_from_txt(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def get_file_size(file_path: str) -> int:
    """Get file size in bytes"""
    return os.path.getsize(file_path)


def is_supported_file_type(filename: str) -> bool:
    """Check if file type is supported"""
    supported_extensions = ['.txt', '.docx', '.pdf']
    _, ext = os.path.splitext(filename)
    return ext.lower() in supported_extensions

# Made with Bob
